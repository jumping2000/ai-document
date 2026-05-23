"""
Export Skill

Responsibilities:
- Convert markdown content to .docx (python-docx)
- Convert markdown content to .pdf (reportlab)
- Save files to configured storage path
- Return file paths

Input:  {content: str, title: str, doc_type: str, workflow_id: str}
Output: {docx_path: str, pdf_path: str}
"""

from __future__ import annotations

import re
import uuid
from pathlib import Path
from typing import Any

import structlog
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    HRFlowable,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from app.core.config import settings

log = structlog.get_logger(__name__)


class ExportSkill:
    """
    Document export to docx and pdf.

    Success criteria:
    - File exists at returned path
    - File size > 0
    - No exceptions

    Failure mode:
    - Raises ExportError with detail message
    - Does NOT retry — caller decides
    """

    def __init__(self) -> None:
        self.base_path = Path(settings.documents_base_path)
        self.base_path.mkdir(parents=True, exist_ok=True)

    async def export_docx(
        self,
        content: str,
        title: str,
        workflow_id: str,
        doc_type: str = "capitolato",
    ) -> str:
        """
        Input:  markdown content string
        Output: absolute file path to .docx
        """
        output_dir = self.base_path / workflow_id
        output_dir.mkdir(parents=True, exist_ok=True)
        file_path = output_dir / f"{doc_type}_{uuid.uuid4().hex[:8]}.docx"

        doc = Document()
        self._setup_doc_styles(doc)
        self._write_markdown_to_doc(doc, content, title)
        doc.save(str(file_path))

        log.info("export_docx_done", path=str(file_path), size=file_path.stat().st_size)
        return str(file_path)

    async def export_pdf(
        self,
        content: str,
        title: str,
        workflow_id: str,
        doc_type: str = "capitolato",
    ) -> str:
        """
        Input:  markdown content string
        Output: absolute file path to .pdf
        """
        output_dir = self.base_path / workflow_id
        output_dir.mkdir(parents=True, exist_ok=True)
        file_path = output_dir / f"{doc_type}_{uuid.uuid4().hex[:8]}.pdf"

        self._write_pdf(content, title, str(file_path))

        log.info("export_pdf_done", path=str(file_path), size=file_path.stat().st_size)
        return str(file_path)

    # ── DOCX helpers ──────────────────────────────────────────────────────────

    def _setup_doc_styles(self, doc: Document) -> None:
        """Apply corporate styling to the document."""
        # Normal style
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10)

        # Title style
        title_style = doc.styles["Title"]
        title_style.font.size = Pt(18)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

    def _write_markdown_to_doc(self, doc: Document, content: str, title: str) -> None:
        """Parse markdown and write to docx paragraph by paragraph."""
        doc.add_heading(title, level=0)

        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue

            if stripped.startswith("#### "):
                doc.add_heading(stripped[5:], level=4)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith(("- ", "* ")):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(stripped[2:])
            elif re.match(r"^\d+\. ", stripped):
                p = doc.add_paragraph(style="List Number")
                p.add_run(re.sub(r"^\d+\. ", "", stripped))
            elif stripped.startswith("> "):
                p = doc.add_paragraph(stripped[2:])
                p.style = doc.styles["Body Text"]
                p.paragraph_format.left_indent = Inches(0.5)
            else:
                p = doc.add_paragraph()
                self._add_inline_formatting(p, stripped)

    def _add_inline_formatting(self, paragraph: Any, text: str) -> None:
        """Handle **bold** and *italic* inline markdown."""
        parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)

    # ── PDF helpers ───────────────────────────────────────────────────────────

    def _write_pdf(self, content: str, title: str, output_path: str) -> None:
        """Build a clean A4 PDF from markdown content."""
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            leftMargin=2.5 * cm,
            rightMargin=2.5 * cm,
            topMargin=2.5 * cm,
            bottomMargin=2.5 * cm,
        )

        styles = getSampleStyleSheet()
        brand_blue = colors.HexColor("#1A1A2E")

        custom_styles = {
            "Title": ParagraphStyle("CustomTitle", parent=styles["Title"],
                                    fontSize=20, textColor=brand_blue, spaceAfter=20),
            "H1": ParagraphStyle("H1", parent=styles["Heading1"],
                                 fontSize=14, textColor=brand_blue, spaceBefore=14, spaceAfter=6),
            "H2": ParagraphStyle("H2", parent=styles["Heading2"],
                                 fontSize=12, textColor=brand_blue, spaceBefore=10, spaceAfter=4),
            "H3": ParagraphStyle("H3", parent=styles["Heading3"],
                                 fontSize=10, spaceBefore=8, spaceAfter=2),
            "Body": ParagraphStyle("Body", parent=styles["Normal"],
                                   fontSize=9, leading=14, spaceAfter=6),
            "Bullet": ParagraphStyle("Bullet", parent=styles["Normal"],
                                     fontSize=9, leftIndent=20, bulletIndent=10, spaceAfter=3),
        }

        story = [Paragraph(title, custom_styles["Title"]), Spacer(1, 0.5 * cm)]

        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 0.2 * cm))
                continue

            if stripped.startswith("## "):
                story.append(HRFlowable(width="100%", thickness=0.5, color=brand_blue))
                story.append(Paragraph(stripped[3:], custom_styles["H1"]))
            elif stripped.startswith("### "):
                story.append(Paragraph(stripped[4:], custom_styles["H2"]))
            elif stripped.startswith("#### "):
                story.append(Paragraph(stripped[5:], custom_styles["H3"]))
            elif stripped.startswith(("- ", "* ")):
                story.append(Paragraph(f"• {stripped[2:]}", custom_styles["Bullet"]))
            else:
                story.append(Paragraph(stripped, custom_styles["Body"]))

        doc.build(story)
